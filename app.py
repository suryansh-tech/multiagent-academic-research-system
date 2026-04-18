import streamlit as st
import time
import io
import ast
import os
import fitz  # PyMuPDF
from markdown_pdf import MarkdownPdf, Section
from docx import Document
from agents import build_reader_agent, build_search_agent, writer_chain, critic_chain, revision_writer_chain, qa_chain, tutor_chain, fact_checker_chain

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResearchMind · AI Research Agent",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=JetBrains+Mono:wght@400;500;700&display=swap');

/* ── Reset & base ── */
html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
    color: #000000;
}

.stApp {
    background-color: #ffffff;
    background-size: 20px 20px;
    background-image: 
        linear-gradient(to right, #f0f0f0 1px, transparent 1px),
        linear-gradient(to bottom, #f0f0f0 1px, transparent 1px);
}
* { border-radius: 0 !important; box-shadow: none !important; }

#MainMenu, footer, header { visibility: hidden; }
.block-container { padding: 2rem 3rem 4rem; max-width: 1000px; }

/* Overrides */
.stTextInput { padding-bottom: 0.8rem !important; }
div[data-testid="stTextInput"] div[data-baseweb="input"], 
div[data-testid="stSelectbox"] div[data-baseweb="select"] { 
    background-color: #f5f5f5 !important; 
    border: none !important; 
    border-bottom: 1px solid #d4d4d4 !important; 
    border-radius: 0 !important; 
    box-shadow: none !important;
}
div[data-testid="stTextInput"] input, 
div[data-testid="stSelectbox"] div[class*="content"] { 
    font-size: 1rem !important; 
    font-weight: 300 !important; 
    color: #000000 !important; 
    padding-left: 0 !important;
}
div[data-testid="stTextInput"] div[data-baseweb="input"]:focus-within, 
div[data-testid="stSelectbox"] div[data-baseweb="select"]:focus-within { 
    border-bottom: 1px solid #000000 !important; 
}
label[data-testid="stWidgetLabel"] p {
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 9px !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
    color: #666666 !important;
    margin-bottom: 0.2rem !important;
}
div[data-testid="stFileUploader"] { padding: 0 !important; }
div[data-testid="stFileUploader"] > section { 
    font-family: 'JetBrains Mono', monospace !important; 
    border: 1px dashed #d4d4d4 !important; 
    background: #f9f9f9 !important;
    padding: 1rem !important;
}
div[data-baseweb="select"] span { padding-left: 0 !important; }

.stButton > button { background: #000000 !important; color: #ffffff !important; font-family: 'JetBrains Mono', monospace !important; font-size: 10px !important; font-weight: 700 !important; text-transform: uppercase !important; letter-spacing: 0.1em !important; border: 1px solid #000000 !important; padding: 0.8rem 1.5rem !important; transition: background 0.2s !important; width: 100%; border-radius: 0 !important; box-shadow: none !important; }
.stButton > button:hover { background: #333333 !important; }

/* Panels and Hero */
.continuous-frame { border: 1px solid #d4d4d4; position: relative; background: #ffffff; padding: 2.5rem; margin-bottom: 2rem; overflow: visible; }
.continuous-frame::before { content: ''; position: absolute; top: -100vh; left: -1px; bottom: -100vh; width: 1px; background: #e5e5e5; z-index: -1; }
.continuous-frame::after { content: ''; position: absolute; top: -1px; left: -100vw; right: -100vw; height: 1px; background: #e5e5e5; z-index: -1; }

.hero-meta { display: flex; justify-content: space-between; align-items: start; margin-bottom: 0.5rem; }
.hero-meta span { font-family: 'JetBrains Mono', monospace; font-size: 10px; text-transform: uppercase; letter-spacing: 0.3em; color: #444444; }
.hero h1 { font-size: 4rem; font-weight: 900; line-height: 0.9; letter-spacing: -0.05em; text-transform: uppercase; color: #000000; margin: 0 0 1.5rem; }
.hero p { font-size: 1rem; color: #444444; font-weight: 300; max-width: 600px; line-height: 1.6; }

.section-heading { font-family: 'JetBrains Mono', monospace; font-size: 10px; letter-spacing: 0.2em; text-transform: uppercase; color: #000; margin-bottom: 1rem; border-bottom: 1px solid #000; padding-bottom: 0.25rem; display: inline-block; font-weight: 700; }
.step-grid { display: grid; grid-template-columns: repeat(2, 1fr); gap: 1px; background: #e5e5e5; border: 1px solid #d4d4d4; }
.step-card { background: #ffffff; padding: 1.25rem; display: flex; flex-direction: column; gap: 0.8rem; margin: 0; border: none; }
.step-card.active { background: #fafafa; position: relative; z-index: 10; box-shadow: inset 0 0 0 1px #000000; }
.step-card.done { background: #ffffff; }
.step-card.skipped { opacity: 0.4; }
.step-num { font-family: 'JetBrains Mono', monospace; font-size: 9px; color: #888888; margin-bottom: 0.2rem; display: block; }
.step-header-row { display: flex; align-items: center; justify-content: space-between; }
.step-title { font-size: 12px; font-weight: 800; text-transform: uppercase; letter-spacing: -0.02em; color: #000; }
.step-status { font-family: 'JetBrains Mono', monospace; font-size: 9px; padding: 2px 6px; border: 1px solid #e5e5e5; }
.status-running { border-color: #000; background: #000; color: #fff; }
.status-done { border-color: #000; color: #000; }
.status-waiting { border-color: #e5e5e5; color: #888888; }
.status-skipped { border-color: transparent; color: #888888; }

.report-panel { background: #ffffff; border: 1px solid #d4d4d4; padding: 3rem 4rem; position: relative; margin-top: 1rem; margin-bottom: 2rem; }
.report-panel::before { content: 'Ref: 882-CR-2024 // MOD: SIGMA'; position: absolute; top: 1rem; left: 1rem; font-family: 'JetBrains Mono', monospace; font-size: 9px; color: #888888; letter-spacing: 0.1em; }
.report-panel h1, .report-panel h2, .report-panel h3 { font-weight: 800; letter-spacing: -0.03em; color: #000; line-height: 1.1; margin-top: 2rem; margin-bottom: 1rem; }
.report-panel h1 { font-size: 2.5rem; }
.report-panel h2 { font-size: 2rem; }
.report-panel h3 { font-size: 14px; font-family: 'JetBrains Mono', monospace; text-transform: uppercase; letter-spacing: 0.1em; border-bottom: 2px solid #000; display: inline-block; padding-bottom: 0.2rem; }
.report-panel p, .report-panel li { font-size: 1.125rem; font-weight: 300; line-height: 1.7; color: #444444; }
.report-panel mark, mark { background: #fffbeb; padding: 0.1rem 0.3rem; border: 1px solid #e5e5e5; color: #000; }
.panel-label { font-family: 'JetBrains Mono', monospace; font-size: 9px; letter-spacing: 0.2em; text-transform: uppercase; color: #444444; margin-bottom: 1.5rem; }

div[data-testid="stChatMessage"] { border: 1px solid #e5e5e5; background: #fafafa; padding: 1.5rem; margin-bottom: 1rem; }
div[data-testid="stChatMessage"][data-testid="chat-message-user"] { background: #000000; color: #ffffff; }
div[data-testid="stChatMessage"][data-testid="chat-message-user"] p, div[data-testid="stChatMessage"][data-testid="chat-message-user"] span { color: #ffffff !important; }
div[data-testid="stChatMessage"] p { font-size: 0.875rem; font-weight: 300; line-height: 1.6; }

div.stChatInputContainer {
    border-radius: 0 !important;
    border: 1px solid #d4d4d4 !important;
    background-color: #f5f5f5 !important;
}

div.stChatFloatingInputContainer {
    max-width: 1000px !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    margin: 0 !important;
    background: transparent !important;
    padding-bottom: 2rem !important;
}
</style>
""", unsafe_allow_html=True)

# ── Session State ─────────────────────────────────────────────────────────────
for key in ("results", "running", "done", "chat_history"):
    if key not in st.session_state:
        st.session_state[key] = {} if key == "results" else ([] if key == "chat_history" else False)

# ── Document Extractors & Generators ──────────────────────────────────────────
def extract_pdf_bytes(pdf_bytes, max_pages=10):
    try:
        doc = fitz.open("pdf", pdf_bytes)
        text = ""
        for i in range(min(len(doc), max_pages)):
            text += doc[i].get_text() + "\n"
        return text
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"

def generate_pdf(text):
    try:
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(text))
        temp_path = "temp_report.pdf"
        pdf.save(temp_path)
        with open(temp_path, "rb") as f:
            data = f.read()
        os.remove(temp_path)
        return data
    except Exception as e:
        return None

def generate_docx(text):
    doc = Document()
    for line in text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def extract_text(c):
    if isinstance(c, list):
        return "\n".join([i.get("text", "") if isinstance(i, dict) else str(i) for i in c])
    if isinstance(c, str):
        try:
            if c.strip().startswith("[{") and c.strip().endswith("]"):
                parsed = ast.literal_eval(c)
                if isinstance(parsed, list):
                    return "\n".join([i.get("text", "") if isinstance(i, dict) else str(i) for i in parsed])
        except Exception:
            pass
    return str(c)

# ── Layout: Hero & Input ──────────────────────────────────────────────────────
st.markdown("""
<div class="continuous-frame hero">
    <div class="hero-meta">
        <span>System Initialized // v4.0.2</span>
        <span>Mode: HI-DENSITY</span>
    </div>
    <h1>Research<br/>Pipeline</h1>
    <p>A high-density collaborative intelligence environment for multi-stage synthesis and technical reporting. Engineered for precision and modular scalability.</p>
</div>
""", unsafe_allow_html=True)

col_input, col_spacer, col_pipeline = st.columns([5, 0.5, 4])

with col_input:
    topic = st.text_input("Research Topic", placeholder="e.g. Quantum computing in 2025", key="topic")
    uploaded_pdf = st.file_uploader("Upload PDF Reference List (Optional)", type=["pdf"], help="Reads up to the first 10 pages locally.")
    depth = st.selectbox("Research Depth", [
        "Quick (Search & Write)", 
        "Standard (Search, Read, Write, Critic)", 
        "Deep Dive (Standard + Revision + Fact Check)",
        "Socratic Tutor (Outline & Teach)"
    ], key="depth")
    run_btn = st.button("⚡  Run Research Pipeline", use_container_width=True)

with col_pipeline:
    st.markdown('<div class="section-heading">Pipeline Status</div>', unsafe_allow_html=True)
    r = st.session_state.results
    
    def get_status(step, required_for_depth=True):
        if not required_for_depth:
            return ("- SKIPPED", "status-skipped", "skipped")
        if not r: return ("WAITING", "status-waiting", "")
        if step in r: return ("✓ DONE", "status-done", "done")
        
        if st.session_state.running:
            keys = {"search":0, "reader":1, "main":2, "critic":3, "revision":4, "fact_checker":5}
            if not any(k in r for k in keys if keys[k] > keys.get(step, 99)):
                return ("● RUNNING", "status-running", "active")
        return ("WAITING", "status-waiting", "")

    is_quick = depth.startswith("Quick")
    is_deep = depth.startswith("Deep")
    is_tutor = depth.startswith("Socratic")

    def render_step(num, title, state_key, req):
        label, cls, card_cls = get_status(state_key, req)
        return f'<div class="step-card {card_cls}"><span class="step-num">0{num} / {state_key.upper()}</span><div class="step-header-row"><span class="step-title">{title}</span><span class="step-status {cls}">{label}</span></div></div>'

    steps_html = '<div class="step-grid">'
    steps_html += render_step("1", "Search Agent", "search", True)
    steps_html += render_step("2", "Reader Agent", "reader", not is_quick and not is_tutor)
    steps_html += render_step("3", "Writer / Tutor", "main", True)
    steps_html += render_step("4", "Critic Chain", "critic", not is_quick and not is_tutor)
    steps_html += render_step("5", "Revision Chain", "revision", is_deep)
    steps_html += render_step("6", "Fact Checker", "fact_checker", is_deep)
    steps_html += '</div>'
    
    st.markdown(steps_html, unsafe_allow_html=True)


# ── Pipeline Execution ────────────────────────────────────────────────────────
if run_btn:
    if not topic.strip():
        st.warning("Please enter a research topic first.")
    else:
        st.session_state.results = {}
        st.session_state.chat_history = []
        if uploaded_pdf is not None:
            st.session_state.pdf_content = extract_pdf_bytes(uploaded_pdf.read())
        else:
            st.session_state.pdf_content = ""
            
        st.session_state.running = True
        st.session_state.done = False
        st.rerun()

if st.session_state.running and not st.session_state.done:
    r = {}
    is_quick = st.session_state.depth.startswith("Quick")
    is_deep = st.session_state.depth.startswith("Deep")
    is_tutor = st.session_state.depth.startswith("Socratic")
    t = st.session_state.topic
    pdf_text = st.session_state.get("pdf_content", "")

    # Step 1
    with st.status("🔍 Search Agent is gathering data...", expanded=True) as status:
        search_agent = build_search_agent()
        sr = search_agent.invoke({"messages": [("user", f"Find detailed information about: {t}")]})
        r["search"] = extract_text(sr["messages"][-1].content)
        st.session_state.results = dict(r)
        status.update(label="🔍 Search Agent completed!", state="complete", expanded=False)

    # Step 2
    if not is_quick and not is_tutor:
        with st.status("📄 Reader Agent is extracting deep content...", expanded=True) as status:
            reader_agent = build_reader_agent()
            rr = reader_agent.invoke({"messages": [("user", f"Pick the most relevant URL and scrape it:\n{r['search'][:800]}")]})
            r["reader"] = extract_text(rr["messages"][-1].content)
            st.session_state.results = dict(r)
            status.update(label="📄 Reader Agent completed!", state="complete", expanded=False)
    else:
        r["reader"] = "Skipped"

    # Step 3
    research_combined = f"SEARCH RESULTS:\n{r['search']}\n\nSCRAPED:\n{r.get('reader','')}"
    if pdf_text:
        research_combined = f"UPLOADED PDF DOCUMENTATION:\n{pdf_text}\n\n" + research_combined
        
    if is_tutor:
        with st.status("👨‍🏫 Socratic Tutor is outlining...", expanded=True) as status:
            stream_container = st.empty()
            full_text = ""
            for chunk in tutor_chain.stream({"topic": t, "research": research_combined}):
                full_text += chunk
                stream_container.markdown(full_text + "▌")
            stream_container.markdown(full_text)
            r["main"] = full_text
            st.session_state.results = dict(r)
            status.update(label="👨‍🏫 Tutor completed!", state="complete", expanded=False)
    else:
        with st.status("✍️ Writer is drafting...", expanded=True) as status:
            stream_container = st.empty()
            full_text = ""
            for chunk in writer_chain.stream({"topic": t, "research": research_combined}):
                full_text += chunk
                stream_container.markdown(full_text + "▌")
            stream_container.markdown(full_text)
            r["main"] = full_text
            st.session_state.results = dict(r)
            status.update(label="✍️ Writer completed!", state="complete", expanded=False)

    # Step 4
    if not is_quick and not is_tutor:
        with st.status("🧐 Critic is reviewing...", expanded=True) as status:
            stream_container = st.empty()
            crit_text = ""
            for chunk in critic_chain.stream({"report": r["main"]}):
                crit_text += chunk
                stream_container.markdown(crit_text + "▌")
            stream_container.markdown(crit_text)
            r["critic"] = crit_text
            st.session_state.results = dict(r)
            status.update(label="🧐 Critic completed!", state="complete", expanded=False)
    else:
        r["critic"] = "Skipped"

    # Step 5
    if is_deep:
        with st.status("🔄 Revision Writer is polishing...", expanded=True) as status:
            stream_container = st.empty()
            full_rev = ""
            for chunk in revision_writer_chain.stream({"topic": t, "draft": r["main"], "feedback": r["critic"]}):
                full_rev += chunk
                stream_container.markdown(full_rev + "▌")
            stream_container.markdown(full_rev)
            r["revision"] = full_rev
            st.session_state.results = dict(r)
            status.update(label="🔄 Revision completed!", state="complete", expanded=False)
            
    # Step 6
    if is_deep:
        with st.status("🛡️ Fact Checker is verifying...", expanded=True) as status:
            stream_container = st.empty()
            fact_rev = ""
            for chunk in fact_checker_chain.stream({"research": research_combined, "draft": r["revision"]}):
                fact_rev += chunk
                stream_container.markdown(fact_rev + "▌")
            stream_container.markdown(fact_rev)
            r["fact_checker"] = fact_rev
            st.session_state.results = dict(r)
            status.update(label="🛡️ Fact Checker completed!", state="complete", expanded=False)
    
    # Finalize
    if "fact_checker" in r:
        r["final_report"] = r["fact_checker"]
    elif "revision" in r:
        r["final_report"] = r["revision"]
    else:
        r["final_report"] = r["main"]
        
    st.session_state.results = dict(r)
    st.session_state.running = False
    st.session_state.done = True
    st.rerun()


# ── Results & Exports ─────────────────────────────────────────────────────────
if st.session_state.done:
    r = st.session_state.results
    
    st.markdown('<div class="report-panel"><div class="panel-label">📝 Final Research Report</div>', unsafe_allow_html=True)
    st.markdown(r["final_report"], unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("### Export Report")
    b1, b2, b3, _ = st.columns([2,2,2,4])
    with b1:
        st.download_button("⬇ Markdown (.md)", r["final_report"], file_name="report.md", mime="text/markdown", use_container_width=True)
    with b2:
        pdf_data = generate_pdf(r["final_report"])
        if pdf_data:
            st.download_button("⬇ PDF (.pdf)", pdf_data, file_name="report.pdf", mime="application/pdf", use_container_width=True)
    with b3:
        docx_data = generate_docx(r["final_report"])
        if docx_data:
            st.download_button("⬇ Word (.docx)", docx_data, file_name="report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    # Follow-up Chat
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.subheader("💬 Ask Follow-up Questions")
    
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"], unsafe_allow_html=True)
            
    if prompt := st.chat_input("Ask something about the report..."):
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
            
        with st.chat_message("assistant"):
            stream_container = st.empty()
            ans = ""
            for chunk in qa_chain.stream({"report": r["final_report"], "question": prompt}):
                ans += chunk
                stream_container.markdown(ans + "▌")
            stream_container.markdown(ans, unsafe_allow_html=True)
        st.session_state.chat_history.append({"role": "assistant", "content": ans})